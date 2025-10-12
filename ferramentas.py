import os
import pandas as pd
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import re

# Processamento de arquivos
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
import openpyxl

# RAG
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document

# Visualizações
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from io import BytesIO
import base64

# CrewAI
from crewai.tools import tool
from dotenv import load_dotenv

load_dotenv()

# ============================================================================
# FERRAMENTA 1: LEITOR UNIVERSAL DE DOCUMENTOS (PDF/WORD/EXCEL)
# ============================================================================

class DocumentReader:
    """Leitor universal de documentos com busca semântica"""
    
    def __init__(self, docs_directory: str = "./documentos"):
        self.docs_directory = Path(docs_directory)
        self.vector_store = None
        self.documents_metadata = {}
        self._initialize()
    
    def _initialize(self):
        """Indexa todos os documentos"""
        
        if not self.docs_directory.exists():
            self.docs_directory.mkdir(parents=True, exist_ok=True)
            print(f"[Doc Reader] Diretório criado: {self.docs_directory}")
            return
        
        all_documents = []
        
        # PDFs
        for pdf_file in self.docs_directory.glob("*.pdf"):
            try:
                docs = self._extract_from_pdf(pdf_file)
                all_documents.extend(docs)
                self.documents_metadata[pdf_file.name] = {
                    'type': 'pdf',
                    'chunks': len(docs),
                    'processed': True
                }
                print(f"[Doc Reader] ✅ PDF processado: {pdf_file.name}")
            except Exception as e:
                print(f"[Doc Reader] ❌ Erro em {pdf_file.name}: {e}")
        
        # Word
        for word_file in self.docs_directory.glob("*.docx"):
            try:
                docs = self._extract_from_word(word_file)
                all_documents.extend(docs)
                self.documents_metadata[word_file.name] = {
                    'type': 'word',
                    'chunks': len(docs),
                    'processed': True
                }
                print(f"[Doc Reader] ✅ Word processado: {word_file.name}")
            except Exception as e:
                print(f"[Doc Reader] ❌ Erro em {word_file.name}: {e}")
        
        # Excel
        for excel_file in self.docs_directory.glob("*.xlsx"):
            try:
                docs = self._extract_from_excel(excel_file)
                all_documents.extend(docs)
                self.documents_metadata[excel_file.name] = {
                    'type': 'excel',
                    'chunks': len(docs),
                    'processed': True
                }
                print(f"[Doc Reader] ✅ Excel processado: {excel_file.name}")
            except Exception as e:
                print(f"[Doc Reader] ❌ Erro em {excel_file.name}: {e}")
        
        if all_documents:
            embeddings = OpenAIEmbeddings(api_key=os.getenv('OPENAI_API_KEY'))
            self.vector_store = FAISS.from_documents(all_documents, embeddings)
            print(f"[Doc Reader] 🎉 Indexados {len(all_documents)} chunks de {len(self.documents_metadata)} documentos")
    
    def _extract_from_pdf(self, pdf_path: Path) -> List[Document]:
        """Extrai texto de PDF"""
        full_text = ""
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    full_text += f"\n--- Página {page_num} ---\n{text}"
        
        return self._chunk_text(full_text, pdf_path.name)
    
    def _extract_from_word(self, word_path: Path) -> List[Document]:
        """Extrai texto de Word"""
        doc = DocxDocument(word_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        return self._chunk_text(full_text, word_path.name)
    
    def _extract_from_excel(self, excel_path: Path) -> List[Document]:
        """Extrai dados de Excel"""
        wb = openpyxl.load_workbook(excel_path)
        full_text = ""
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            full_text += f"\n\n=== Planilha: {sheet_name} ===\n\n"
            
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                if row_text.strip():
                    full_text += row_text + "\n"
        
        return self._chunk_text(full_text, excel_path.name)
    
    def _chunk_text(self, text: str, source: str) -> List[Document]:
        """Divide texto em chunks"""
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        chunks = splitter.split_text(text)
        
        return [
            Document(
                page_content=chunk,
                metadata={'source': source, 'chunk_id': i}
            )
            for i, chunk in enumerate(chunks)
        ]
    
    def search(self, query: str, k: int = 3) -> str:
        """Busca informações nos documentos"""
        
        if not self.vector_store:
            return "❌ Nenhum documento indexado. Adicione arquivos em ./documentos/"
        
        docs = self.vector_store.similarity_search(query, k=k)
        
        if not docs:
            return f"❌ Nenhuma informação encontrada sobre: {query}"
        
        response = f"📚 INFORMAÇÕES DOS DOCUMENTOS:\n\n"
        for i, doc in enumerate(docs, 1):
            source = doc.metadata.get('source', 'Desconhecido')
            response += f"[{i}] 📄 {source}\n"
            response += f"{doc.page_content}\n"
            response += "-" * 80 + "\n\n"
        
        return response


# ============================================================================
# FERRAMENTA 2: ANALISADOR DE CSV E BANCO LOCAL
# ============================================================================

class CSVDatabaseAnalyzer:
    """Analisador de dados CSV como banco de dados local"""
    
    def __init__(self, csv_directory: str = "./dados"):
        self.csv_directory = Path(csv_directory)
        self.dataframes = {}
        self._load_csvs()
    
    def _load_csvs(self):
        """Carrega todos os CSVs"""
        
        if not self.csv_directory.exists():
            self.csv_directory.mkdir(parents=True, exist_ok=True)
            print(f"[CSV Analyzer] Diretório criado: {self.csv_directory}")
            return
        
        for csv_file in self.csv_directory.glob("*.csv"):
            try:
                df = pd.read_csv(csv_file)
                table_name = csv_file.stem
                self.dataframes[table_name] = df
                print(f"[CSV Analyzer] ✅ Carregado: {table_name} ({len(df)} linhas)")
            except Exception as e:
                print(f"[CSV Analyzer] ❌ Erro em {csv_file.name}: {e}")
    
    def query(self, question: str) -> str:
        """Responde perguntas sobre os dados"""
        
        if not self.dataframes:
            return "❌ Nenhum arquivo CSV encontrado em ./dados/"
        
        # Listar tabelas disponíveis
        tables_info = "\n".join([
            f"- {name}: {len(df)} linhas, {len(df.columns)} colunas"
            for name, df in self.dataframes.items()
        ])
        
        response = f"📊 DADOS DISPONÍVEIS:\n{tables_info}\n\n"
        
        # Análise simples baseada em palavras-chave
        question_lower = question.lower()
        
        # Exemplo: "quantos pacientes"
        if "pacientes" in question_lower and "quantos" in question_lower:
            if "pacientes" in self.dataframes:
                total = len(self.dataframes["pacientes"])
                response += f"Total de pacientes: {total}\n"
        
        # Exemplo: "média de idade"
        if "média" in question_lower or "media" in question_lower:
            if "pacientes" in self.dataframes:
                df = self.dataframes["pacientes"]
                if "idade" in df.columns:
                    media = df["idade"].mean()
                    response += f"Média de idade: {media:.1f} anos\n"
        
        # Exemplo: "total de consultas"
        if "consultas" in question_lower and "total" in question_lower:
            if "consultas" in self.dataframes:
                total = len(self.dataframes["consultas"])
                response += f"Total de consultas: {total}\n"
        
        # Estatísticas gerais
        response += "\n📈 ESTATÍSTICAS:\n"
        for name, df in self.dataframes.items():
            response += f"\n{name.upper()}:\n"
            response += f"  Registros: {len(df)}\n"
            response += f"  Colunas: {', '.join(df.columns.tolist()[:5])}\n"
            
            # Mostrar primeiras linhas
            response += f"\nPrimeiras linhas de {name}:\n"
            response += df.head(3).to_string(index=False) + "\n"
        
        return response
    
    def get_statistics(self, table_name: str) -> str:
        """Retorna estatísticas de uma tabela"""
        
        if table_name not in self.dataframes:
            available = ", ".join(self.dataframes.keys())
            return f"❌ Tabela '{table_name}' não encontrada. Disponíveis: {available}"
        
        df = self.dataframes[table_name]
        
        response = f"📊 ESTATÍSTICAS: {table_name}\n\n"
        response += f"Total de registros: {len(df)}\n"
        response += f"Colunas: {len(df.columns)}\n\n"
        
        # Estatísticas numéricas
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            response += "COLUNAS NUMÉRICAS:\n"
            response += df[numeric_cols].describe().to_string() + "\n\n"
        
        # Valores únicos
        response += "VALORES ÚNICOS:\n"
        for col in df.columns[:5]:
            unique = df[col].nunique()
            response += f"  {col}: {unique} valores únicos\n"
        
        return response


# ============================================================================
# FERRAMENTA 3: GERADOR DE VISUALIZAÇÕES
# ============================================================================

class VisualizationGenerator:
    """Gera gráficos e visualizações"""
    
    def __init__(self, output_dir: str = "./visualizacoes"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def create_chart(self, data_dict: Dict[str, Any], chart_type: str = "bar") -> str:
        """Cria gráfico e retorna path"""
        
        try:
            # Dados de exemplo se não fornecidos
            if not data_dict:
                data_dict = {
                    'labels': ['Jan', 'Fev', 'Mar', 'Abr', 'Mai'],
                    'values': [65, 72, 68, 75, 70]
                }
            
            labels = data_dict.get('labels', [])
            values = data_dict.get('values', [])
            
            # Criar figura
            fig = go.Figure()
            
            if chart_type == "bar":
                fig.add_trace(go.Bar(x=labels, y=values))
            elif chart_type == "line":
                fig.add_trace(go.Scatter(x=labels, y=values, mode='lines+markers'))
            elif chart_type == "pie":
                fig.add_trace(go.Pie(labels=labels, values=values))
            
            fig.update_layout(
                title=data_dict.get('title', 'Gráfico'),
                template='plotly_white'
            )
            
            # Salvar
            filename = f"chart_{chart_type}_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.html"
            filepath = self.output_dir / filename
            fig.write_html(str(filepath))
            
            return f"✅ Gráfico criado: {filepath}\nAcesse o arquivo para visualizar."
            
        except Exception as e:
            return f"❌ Erro ao criar gráfico: {str(e)}"


# ============================================================================
# FERRAMENTA 4: BUSCADOR WEB PARA GUIDELINES MÉDICAS
# ============================================================================

@tool("buscar_guidelines_medicas")
def buscar_guidelines_medicas(query: str) -> str:
    """
    Busca guidelines e protocolos médicos na web.
    Útil para encontrar diretrizes clínicas, protocolos de tratamento,
    recomendações de sociedades médicas, etc.
    
    Args:
        query: Termo de busca (ex: "protocolo diabetes tipo 2")
        
    Returns:
        Informações relevantes de fontes médicas confiáveis
    """
    
    # Simulação - em produção, integraria com APIs reais
    guidelines = {
        "diabetes": "Diretrizes SBD 2024: HbA1c alvo <7%, monitoramento trimestral, metformina primeira linha.",
        "hipertensão": "Diretriz Brasileira 2020: PA alvo <140/90mmHg, IECA/BRA primeira escolha.",
        "covid": "Protocolo MS 2024: Isolamento 5 dias, sintomáticos tratamento suportivo.",
    }
    
    query_lower = query.lower()
    
    response = "🔍 GUIDELINES MÉDICAS:\n\n"
    
    found = False
    for key, value in guidelines.items():
        if key in query_lower:
            response += f"📋 {key.upper()}:\n{value}\n\n"
            found = True
    
    if not found:
        response += f"⚠️ Nenhuma guideline específica encontrada para '{query}'.\n"
        response += "Recomendo consultar:\n"
        response += "- Sociedade Brasileira de Medicina\n"
        response += "- Ministério da Saúde\n"
        response += "- UpToDate / DynaMed\n"
    
    response += "\n⚠️ IMPORTANTE: Sempre verificar fontes primárias e validade das diretrizes."
    
    return response


# ============================================================================
# INSTANCIAR FERRAMENTAS
# ============================================================================

doc_reader = DocumentReader()
csv_analyzer = CSVDatabaseAnalyzer()
viz_generator = VisualizationGenerator()


# ============================================================================
# DECORATORS PARA CREWAI
# ============================================================================

@tool("ler_documentos")
def ler_documentos(pergunta: str) -> str:
    """
    Busca informações em documentos (PDF, Word, Excel).
    Use para encontrar protocolos, relatórios, diretrizes documentadas.
    
    Args:
        pergunta: O que você quer saber dos documentos
        
    Returns:
        Informações encontradas com fonte
    """
    return doc_reader.search(pergunta)


@tool("consultar_banco_dados")
def consultar_banco_dados(pergunta: str) -> str:
    """
    Consulta dados em arquivos CSV (pacientes, consultas, custos, etc).
    Use para análises quantitativas, estatísticas, contagens.
    
    Args:
        pergunta: Pergunta sobre os dados (ex: "quantos pacientes diabéticos?")
        
    Returns:
        Análise dos dados com números e estatísticas
    """
    return csv_analyzer.query(pergunta)


@tool("gerar_visualizacao")
def gerar_visualizacao(descricao: str) -> str:
    """
    Cria gráficos e visualizações.
    Use quando precisar ilustrar dados visualmente.
    
    Args:
        descricao: Tipo de gráfico e dados (ex: "gráfico de barras com consultas por mês")
        
    Returns:
        Path do arquivo gerado
    """
    # Exemplo simples - em produção, parsearia a descrição
    return viz_generator.create_chart({}, "bar")